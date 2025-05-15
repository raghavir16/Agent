from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

class ProposalDocumentGenerator:
    def __init__(self):
        self.doc = Document()
        self.current_section = None

    def create_title_page(self, title, company_name, date):
        """Create the title page of the proposal"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.font.size = Pt(24)
        run.font.bold = True
        
        self.doc.add_paragraph()  # Spacing
        
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(f"Prepared for: {company_name}").font.size = Pt(14)
        
        self.doc.add_paragraph()  # Spacing
        
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(date).font.size = Pt(12)
        
        self.doc.add_page_break()

    def add_section(self, title, content):
        """Add a new section to the document"""
        self.current_section = self.doc.add_heading(title, level=1)
        if isinstance(content, str):
            self.doc.add_paragraph(content)
        elif isinstance(content, list):
            for item in content:
                self.doc.add_paragraph(f"• {item}", style='List Bullet')
        elif isinstance(content, dict):
            for key, value in content.items():
                self.doc.add_heading(key, level=2)
                if isinstance(value, list):
                    for item in value:
                        self.doc.add_paragraph(f"• {item}", style='List Bullet')
                else:
                    self.doc.add_paragraph(str(value))

    def add_table(self, headers, rows):
        """Add a table to the document"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)

    def add_costs_section(self, costs_data):
        """Add the costs section with tables"""
        self.add_section("Costs", "")
        
        # Resource costs table
        if costs_data.get('resource_costs'):
            self.doc.add_heading("Resource Costs", level=2)
            headers = ["Resource", "Quantity", "Unit Cost", "Total"]
            rows = [[cost['resource'], cost['quantity'], f"${cost['unit_cost']}", f"${cost['total']}"] 
                   for cost in costs_data['resource_costs']]
            self.add_table(headers, rows)
        
        # License costs table
        if costs_data.get('license_costs'):
            self.doc.add_heading("License Costs", level=2)
            headers = ["License", "Type", "Cost"]
            rows = [[cost['name'], cost['type'], f"${cost['amount']}"] 
                   for cost in costs_data['license_costs']]
            self.add_table(headers, rows)

    def add_raid_section(self, raid_data):
        """Add the RAID section"""
        self.add_section("RAID Analysis", "")
        
        # Risks
        if raid_data.get('risks'):
            self.doc.add_heading("Risks", level=2)
            headers = ["Risk", "Impact", "Probability", "Mitigation"]
            rows = [[risk['description'], risk['impact'], risk['probability'], risk['mitigation']] 
                   for risk in raid_data['risks']]
            self.add_table(headers, rows)
        
        # Assumptions
        if raid_data.get('assumptions'):
            self.add_section("Assumptions", raid_data['assumptions'])
        
        # Dependencies
        if raid_data.get('dependencies'):
            self.add_section("Dependencies", raid_data['dependencies'])

    def save(self, filename="proposal.docx"):
        """Save the document"""
        # Create 'output' directory if it doesn't exist
        os.makedirs('output', exist_ok=True)
        filepath = os.path.join('output', filename)
        self.doc.save(filepath)
        return filepath 